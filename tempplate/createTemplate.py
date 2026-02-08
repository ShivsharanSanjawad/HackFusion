from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def shade_cell(cell, color_hex):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell border with specified color and width"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '12')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def create_template():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- 1. TITLE & HEADER ---
    title = doc.add_heading('OFFICIAL CIVIL ISSUE REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Color title professionally
    title_format = title.runs[0]
    title_format.font.color.rgb = RGBColor(31, 78, 120)  # Professional dark blue
    title_format.font.size = Pt(28)
    title_format.font.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Civic Issue Tracking & Resolution System')
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray

    # Add generation date placeholder
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run('Generated: ${.now?string("yyyy-MM-dd HH:mm:ss")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(128, 128, 128)  # Dark gray
    
    # Add horizontal line
    line_p = doc.add_paragraph()
    line_p_format = line_p.paragraph_format
    pPr = line_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D9D9D9')
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()  # Spacer

    # --- 2. REPORT DETAILS TABLE (2 Columns) ---
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)

    # Helper to add rows with alternating colors
    def add_row(label, value_placeholder, is_header=False, shade=False):
        row_cells = table.add_row().cells
        
        if is_header:
            # Header styling
            shade_cell(row_cells[0], '1F4E78')  # Dark blue
            shade_cell(row_cells[1], '1F4E78')
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                        run.font.size = Pt(11)
        else:
            # Data row styling
            if shade:
                shade_cell(row_cells[0], 'E8F0F7')  # Light blue
                shade_cell(row_cells[1], 'E8F0F7')
            
            # Label Cell
            p_label = row_cells[0].paragraphs[0]
            run_label = p_label.add_run(label)
            run_label.font.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = RGBColor(31, 78, 120)  # Dark blue
            
            # Value Cell
            p_value = row_cells[1].paragraphs[0]
            run_value = p_value.add_run(value_placeholder)
            run_value.font.size = Pt(10)
        
        # Add borders to cells
        for cell in row_cells:
            set_cell_border(cell, color='1F4E78')

    # Header row
    add_row("Field", "Value", is_header=True)
    
    # Data rows with alternating shading
    add_row("Report ID:", "${id}", shade=False)
    add_row("Date Filed:", "${entryDate}", shade=True)
    add_row("Issue Since:", "${issueSince}", shade=False)
    add_row("Department:", "${department}", shade=True)
    add_row("Current Status:", "${status}", shade=False)
    add_row("Priority Level:", "${priority}", shade=True)
    add_row("Reported By:", "${operatorName}", shade=False)
    add_row("Location (Lat/Lon):", "${lat}, ${lon}", shade=True)
    add_row("Community Support:", "${upvotes} upvotes", shade=False)

    doc.add_paragraph()  # Spacer

    # --- 3. DESCRIPTION SECTION ---
    desc_heading = doc.add_heading('Issue Description', level=1)
    desc_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)  # Professional blue
    desc_heading.runs[0].font.size = Pt(14)
    desc_heading.paragraph_format.space_before = Pt(6)
    desc_heading.paragraph_format.space_after = Pt(6)
    
    desc_p = doc.add_paragraph("${description}")
    desc_p.paragraph_format.space_after = Pt(6)
    desc_p.paragraph_format.space_before = Pt(0)
    desc_p.paragraph_format.left_indent = Inches(0.3)
    desc_p.paragraph_format.right_indent = Inches(0.3)
    desc_run = desc_p.runs[0]
    desc_run.font.size = Pt(11)

    doc.add_paragraph()  # Spacer

    # --- 4. STATUS HISTORY TABLE (Dynamic List) ---
    hist_heading = doc.add_heading('Status Timeline', level=1)
    hist_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    hist_heading.runs[0].font.size = Pt(14)
    hist_heading.paragraph_format.space_before = Pt(6)
    hist_heading.paragraph_format.space_after = Pt(6)

    hist_table = doc.add_table(rows=1, cols=2)
    hist_table.style = 'Table Grid'
    hist_table.autofit = False
    
    # Header Row
    hdr_cells = hist_table.rows[0].cells
    hdr_cells[0].width = Inches(2.2)
    hdr_cells[1].width = Inches(4.3)
    
    # Header styling
    shade_cell(hdr_cells[0], '4472C4')  # Slightly lighter blue
    shade_cell(hdr_cells[1], '4472C4')
    
    hdr_cells[0].text = 'Date of Update'
    hdr_cells[1].text = 'Status Change'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                run.font.size = Pt(11)
        set_cell_border(cell, color='1F4E78')

    # Data Row (The Loop) with proper Freemarker syntax
    row_cells = hist_table.add_row().cells
    row_cells[0].width = Inches(2.2)
    row_cells[1].width = Inches(4.3)
    
    # Shade alternating rows
    shade_cell(row_cells[0], 'D9E2F3')  # Light blue background
    shade_cell(row_cells[1], 'D9E2F3')
    
    # Cell 1: Start of loop + Date
    row_cells[0].text = "[#list history as item] ${item.date}"
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Cell 2: Status + End of loop
    row_cells[1].text = "${item.status} [/#list]"
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Add borders
    for cell in row_cells:
        set_cell_border(cell, color='1F4E78')

    doc.add_paragraph()  # Spacer

    # --- 5. FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.text = "Confidential Report - HackFusion System - Civic Issue Tracking Platform"
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = p_foot.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.italic = True

    # Save
    doc.save('template.docx')
    print("✅ template.docx created successfully with enhanced formatting!")
    print("📋 Features:")
    print("   • Professional color scheme (Dark blue headers)")
    print("   • Alternating row colors for better readability")
    print("   • Improved table borders and styling")
    print("   • Better typography and spacing")
    print("   • Subtitle and timestamp")
    print("   • All fields properly formatted for PDF conversion")

if __name__ == "__main__":
    create_template()